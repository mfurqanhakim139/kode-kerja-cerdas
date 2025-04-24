import docx
from pathlib import Path

# Persiapan: Buat file Word contoh jika belum ada
nama_file_word = 'contoh_dokumen.docx'
file_word = Path(nama_file_word)
if not file_word.exists():
    try:
        doc_baru = docx.Document()
        doc_baru.add_heading('Judul Dokumen Contoh', level=1)
        p1 = doc_baru.add_paragraph('Ini adalah paragraf pertama. ')
        p1.add_run('Bagian ini tebal.').bold = True
        p1.add_run(' Bagian ini biasa lagi.')
        doc_baru.add_paragraph('Ini paragraf kedua dengan gaya berbeda.', style='Intense Quote')
        doc_baru.save(nama_file_word)
        print(f"File '{nama_file_word}' berhasil dibuat untuk demo.")
    except Exception as e:
        print(f"Gagal membuat file Word demo: {e}")


# --- Membaca File Word ---
if file_word.exists():
    try:
        doc = docx.Document(nama_file_word)
        print(f"Berhasil membuka: {nama_file_word}")

        print("\nMembaca teks per paragraf:")
        for i, p in enumerate(doc.paragraphs):
            # Jangan cetak paragraf kosong (biasanya ada di antara elemen)
            if p.text.strip():
                print(f"Paragraf {i}: {p.text}")
                # Contoh melihat run dalam paragraf pertama
                # if i == 1: # Paragraf kedua secara teknis (indeks 1)
                #     print("  Runs di dalamnya:")
                #     for j, run in enumerate(p.runs):
                #         print(f"    Run {j}: '{run.text}' (Tebal: {run.bold}, Miring: {run.italic})")

        # Cara mudah mendapatkan seluruh teks (kurang memperhatikan format)
        # teks_lengkap = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        # print("\n--- Teks Lengkap ---")
        # print(teks_lengkap)

    except Exception as e:
        print(f"Gagal membaca file Word: {e}")
else:
    print(f"File Word '{nama_file_word}' tidak ditemukan.")
