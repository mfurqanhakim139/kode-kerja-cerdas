# Perlu install: pip install python-docx
import docx
from docx.shared import Inches # Untuk mengatur ukuran gambar
from pathlib import Path # <-- Tambahkan impor ini

nama_file_baru = 'dokumen_hasil_python.docx'
# --- Pastikan file gambar ini ADA di folder yang sama dengan skrip ---
# --- atau ganti dengan path lengkap ke gambar Anda ---
nama_file_gambar = 'python.jpg' # Ganti jika nama gambar Anda berbeda

try:
    # 1. Buat dokumen baru
    doc = docx.Document()

    # 2. Tambah Judul
    doc.add_heading('Laporan Otomatis Python', level=0) # Judul Utama

    # 3. Tambah Paragraf biasa
    p1 = doc.add_paragraph('Ini adalah laporan yang dibuat secara otomatis menggunakan ')
    # Tambah Run dengan gaya
    run_python = p1.add_run('Python')
    run_python.bold = True
    p1.add_run(' dan library ')
    run_docx = p1.add_run('python-docx')
    run_docx.italic = True
    p1.add_run('.')

    doc.add_heading('Bagian 1: Pendahuluan', level=1)
    doc.add_paragraph('Ini isi bagian pendahuluan...', style='List Bullet') # Contoh style

    doc.add_heading('Bagian 2: Data', level=1)
    doc.add_paragraph('Menambahkan gambar di bawah ini:')

    # 4. Tambah Gambar
    path_gambar = Path(nama_file_gambar)
    if path_gambar.is_file(): # Cek apakah file gambar benar-benar ada
        try:
            # Hapus baris .touch() - kita butuh gambar asli, bukan file kosong
            # Path(logo_path).touch(exist_ok=True)
            print(f"Mencoba menambahkan gambar: {path_gambar}")
            doc.add_picture(str(path_gambar), width=Inches(2.0)) # Gunakan str(path_gambar)
            print("Gambar berhasil ditambahkan.")
        except FileNotFoundError:
            # Seharusnya tidak terjadi karena sudah dicek is_file(), tapi untuk jaga-jaga
            print(f"Error: File gambar '{path_gambar}' tidak ditemukan saat add_picture.")
            doc.add_paragraph(f'(Gambar "{nama_file_gambar}" tidak ditemukan)')
        except Exception as e_img:
            # Tangkap error lain saat mencoba membaca/menyisipkan gambar
            print(f"Error saat menambahkan gambar '{path_gambar}': {e_img}")
            doc.add_paragraph(f'(Error menambahkan gambar: {e_img})')
    else:
        print(f"Peringatan: File gambar '{nama_file_gambar}' tidak ditemukan di direktori skrip.")
        doc.add_paragraph(f'(File gambar "{nama_file_gambar}" tidak ditemukan)')


    # 5. Tambah Page Break (opsional)
    # doc.add_page_break()
    # doc.add_paragraph('Ini Halaman Baru')

    # 6. Simpan dokumen
    doc.save(nama_file_baru)
    print(f"\nDokumen '{nama_file_baru}' berhasil dibuat.")

except ImportError:
     print("Error: Library 'python-docx' belum terinstal.")
     print("Silakan instal dengan: pip install python-docx")
except Exception as e:
    print(f"Gagal membuat dokumen Word: {e}")